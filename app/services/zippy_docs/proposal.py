"""Business Proposal generator — rewrites Beacon's Drive proposal template.

Flow:
  1. inspect_proposal_template() — downloads the Drive template and reports
     the number of rewritable content sections.
  2. generate(data) — downloads template, extracts paragraphs, asks Claude
     to rewrite non-structural content from transcripts + email threads,
     patches rewrites back preserving formatting, uploads to Google Docs.
  3. update(drive_file_id, change_request, data) — re-generates with a
     change request appended to context and overwrites or creates a new
     version in Drive.
"""
from __future__ import annotations

import logging
from dataclasses import dataclass, field
from datetime import datetime
from typing import Optional

from docx import Document
from docx.shared import Pt, RGBColor
from sqlmodel import select as sm_select

from app.clients.anthropic_client import get_anthropic_client
from app.clients.google_drive import upload_as_google_doc
from app.config import settings
from app.database import AsyncSessionLocal as async_session
from app.models.user_email_connection import UserEmailConnection
from app.models.zippy import IndexedDriveFile
from app.services.zippy_docs.base import (
    GeneratedDocument,
    build_output_path,
    cache_upload,
    get_cached_upload,
    human_today,
)
from app.services.zippy_docs.claude_rewriter import rewrite_with_claude
from app.services.zippy_docs.doc_rewriter import (
    extract_docx_structure,
    rewrite_docx_content,
)

logger = logging.getLogger(__name__)

PROPOSAL_TEMPLATE_DRIVE_ID = "1doYGGVl3uDHXtFbovC9vTMPbIqoF3pyjve_MdwOwE_k"


class ProposalTemplateUnavailable(Exception):
    pass


@dataclass
class ProposalInput:
    client_name: str
    variant: str = "main"
    email_thread_content: str = ""
    transcript: str = ""
    prepared_by: str = ""
    prepared_by_title: str = ""
    prepared_by_phone: str = ""
    prepared_by_email: str = ""
    date: str = ""
    platform: str = ""
    domain: str = ""
    client_description: str = ""
    use_cases: list[str] = field(default_factory=list)
    effort_reduction_pct: str = ""
    timeline_reduction_pct: str = ""
    hypercare_reduction_pct: str = ""
    annual_platform_fee: str = ""
    per_client_fee: str = ""
    implementations_per_year: str = ""
    avg_hours_per_impl: str = ""
    hourly_rate: str = ""
    change_request: str = ""


async def _get_connection(user_id: Optional[str]):
    from sqlalchemy import or_, case as sa_case

    async with async_session() as session:
        stmt = sm_select(UserEmailConnection).where(
            UserEmailConnection.is_active == True,  # noqa: E712
        )
        if user_id is not None:
            stmt = stmt.where(
                or_(
                    UserEmailConnection.user_id == user_id,
                    UserEmailConnection.is_admin_folder == True,  # noqa: E712
                )
            )
            priority = sa_case(
                (UserEmailConnection.user_id == user_id, 0), else_=1
            )
            stmt = stmt.order_by(priority)
        result = await session.execute(stmt.limit(1))
        return result.scalar_one_or_none()


async def _export_google_doc_as_docx(file_id: str, connection) -> Optional[bytes]:
    """Export a Google Doc to .docx bytes using Drive's export endpoint."""
    import httpx
    from app.clients.google_drive import _ensure_token, DRIVE_API_BASE

    access_token, _updated = await _ensure_token(
        connection.token_data,
        settings.gmail_client_id,
        settings.gmail_client_secret,
    )
    docx_mime = (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    async with httpx.AsyncClient(timeout=60) as http:
        resp = await http.get(
            f"{DRIVE_API_BASE}/files/{file_id}/export",
            headers={"Authorization": f"Bearer {access_token}"},
            params={"mimeType": docx_mime},
        )
    if resp.status_code == 200:
        return resp.content
    logger.warning(
        "Proposal template export failed (status %s): %s",
        resp.status_code, resp.text[:300],
    )
    return None


async def _fetch_template_bytes(user_id: Optional[str]) -> bytes:
    from app.clients import google_drive

    connection = await _get_connection(user_id)
    if not connection:
        raise ProposalTemplateUnavailable("No active Drive connection")

    # 1. Hardcoded Drive ID — Google Doc export
    try:
        data = await _export_google_doc_as_docx(
            PROPOSAL_TEMPLATE_DRIVE_ID, connection
        )
        if data and data.startswith(b"PK"):
            return data
    except Exception as exc:
        logger.warning("Hardcoded proposal template fetch failed: %s", exc)

    # 2. Indexed Drive search fallback
    _KEYWORDS = [
        "copy of business proposal template",
        "business proposal template",
        "beacon business proposal",
        "proposal template",
    ]
    async with async_session() as session:
        for kw in _KEYWORDS:
            result = await session.execute(
                sm_select(IndexedDriveFile).where(
                    IndexedDriveFile.name.ilike(f"%{kw}%")
                ).limit(1)
            )
            row = result.scalar_one_or_none()
            if row:
                try:
                    if row.mime_type == "application/vnd.google-apps.document":
                        data = await _export_google_doc_as_docx(
                            row.drive_file_id, connection
                        )
                    else:
                        data, _mime, _updated = await google_drive.download_file_bytes(
                            file_id=row.drive_file_id,
                            mime_type=row.mime_type or "",
                            token_data=connection.token_data,
                            client_id=settings.gmail_client_id,
                            client_secret=settings.gmail_client_secret,
                        )
                    if data and data.startswith(b"PK"):
                        return data
                except Exception as exc:
                    logger.warning(
                        "Proposal template fetch via index '%s' failed: %s",
                        row.name, exc,
                    )

    raise ProposalTemplateUnavailable(
        "Business Proposal template not found in Drive"
    )


async def inspect_proposal_template(user_id: Optional[str] = None) -> dict:
    try:
        docx_bytes = await _fetch_template_bytes(user_id)
        structure = extract_docx_structure(docx_bytes)
        content_blocks = [b for b in structure if not b.get("is_structural")]
        return {
            "available": True,
            "template_drive_id": PROPOSAL_TEMPLATE_DRIVE_ID,
            "total_blocks": len(structure),
            "rewritable_blocks": len(content_blocks),
            "message": (
                f"Proposal template found - "
                f"{len(content_blocks)} rewritable content sections ready. "
                f"Template has two variants: 'lite' (7 sections, concise) "
                f"and 'main' (9 sections, full detail). "
                f"Ask the user which they prefer, or default to 'main' for new proposals."
            ),
        }
    except ProposalTemplateUnavailable as exc:
        return {"available": False, "message": str(exc)}
    except Exception as exc:
        return {"available": False, "message": f"Template error: {exc}"}


async def generate(
    data: ProposalInput, user_id: Optional[str] = None
) -> GeneratedDocument:
    path, url = build_output_path("proposal", data.client_name)
    created_at = datetime.utcnow()

    try:
        docx_bytes = await _fetch_template_bytes(user_id)
        structure = extract_docx_structure(docx_bytes)

        user_inputs = {
            "client_name": data.client_name,
            "variant": data.variant or "main",
            "prepared_by": data.prepared_by or "Beacon Team",
            "prepared_by_title": data.prepared_by_title or "",
            "prepared_by_phone": data.prepared_by_phone or "",
            "prepared_by_email": data.prepared_by_email or "",
            "date": data.date or human_today(),
            "platform": data.platform or "",
            "domain": data.domain or "",
            "client_description": data.client_description or "",
            "use_cases": data.use_cases,
            "effort_reduction_pct": data.effort_reduction_pct or "50-60",
            "timeline_reduction_pct": data.timeline_reduction_pct or "40-60",
            "hypercare_reduction_pct": data.hypercare_reduction_pct or "70-80",
            "annual_platform_fee": data.annual_platform_fee or "",
            "per_client_fee": data.per_client_fee or "",
            "implementations_per_year": data.implementations_per_year or "",
            "avg_hours_per_impl": data.avg_hours_per_impl or "",
            "hourly_rate": data.hourly_rate or "",
            "email_thread": data.email_thread_content or "",
            "transcript": data.transcript or "",
            "change_request": data.change_request or "",
        }

        client = get_anthropic_client()
        rewrites = await rewrite_with_claude(
            structure=structure,
            user_inputs=user_inputs,
            doc_type="proposal",
            client=client,
            model=settings.CLAUDE_MODEL_STANDARD,
        )
        out_bytes = rewrite_docx_content(docx_bytes, rewrites)
        path.write_bytes(out_bytes)

    except ProposalTemplateUnavailable:
        logger.warning("Proposal template unavailable - using fallback renderer")
        _render_fallback(data, path)

    result = GeneratedDocument(
        filename=path.name,
        path=str(path),
        url=url,
        kind="proposal_docx",
        summary=f"Business Proposal for {data.client_name}",
        created_at=created_at,
    )
    await _try_upload_to_drive(
        result, path, data.client_name,
        user_id=user_id,
        change_request=data.change_request,
    )
    return result


def _render_fallback(data: ProposalInput, path) -> None:
    doc = Document()
    t = doc.add_heading(f"Business Proposal - {data.client_name}", level=0)
    for run in t.runs:
        run.font.color.rgb = RGBColor(0x17, 0x4A, 0x8B)
    for section in [
        "Executive Summary", "Current Challenges", "Proposed Solution",
        "Expected ROI & Impact", "Implementation Plan", "Commercials", "Next Steps"
    ]:
        h = doc.add_heading(section, level=1)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0x17, 0x4A, 0x8B)
        doc.add_paragraph("[Content to be added]")
    footer = doc.add_paragraph()
    fr = footer.add_run(
        f"\nGenerated by Zippy on "
        f"{datetime.utcnow().strftime('%d %b %Y %H:%M UTC')}"
    )
    fr.italic = True
    fr.font.size = Pt(9)
    fr.font.color.rgb = RGBColor(0x8A, 0x8A, 0x8A)
    doc.save(str(path))


async def _try_upload_to_drive(
    doc: GeneratedDocument,
    path,
    client_name: str,
    *,
    user_id: Optional[str],
    change_request: str = "",
) -> None:
    if doc.drive_url and not change_request:
        return
    cache_key = client_name
    if user_id and not change_request:
        cached = get_cached_upload(str(user_id), cache_key, doc.kind)
        if cached:
            doc.drive_url = cached
            logger.info("Reusing cached proposal Drive upload: %s", cached)
            return
    try:
        connection = await _get_connection(user_id)
        if not connection:
            logger.info("No active Drive connection - skipping upload")
            return

        docx_bytes = path.read_bytes()
        gdoc_name = (
            f"{client_name} - Business Proposal - "
            f"{doc.created_at.strftime('%d %b %Y')}"
        )
        folder_id = connection.selected_drive_folder_id or None

        file_id, web_view_link = await upload_as_google_doc(
            filename=gdoc_name,
            docx_bytes=docx_bytes,
            token_data=connection.token_data,
            client_id=settings.gmail_client_id,
            client_secret=settings.gmail_client_secret,
            parent_folder_id=folder_id,
        )
        doc.drive_file_id = file_id
        doc.drive_url = web_view_link
        logger.info("Proposal uploaded to Google Docs: %s", web_view_link)
        if user_id and web_view_link and not change_request:
            cache_upload(str(user_id), cache_key, doc.kind, web_view_link)
    except PermissionError as exc:
        logger.info("drive.file scope not granted - skipping upload: %s", exc)
    except Exception as exc:
        logger.warning("Google Docs upload failed (non-fatal): %s", exc)
