#!/usr/bin/env python3
"""Generate the Recotap integration proposal as a Word (.docx) file.

Reusable converter for the partner-facing proposal: edit the BLOCKS list (or the
companion RECOTAP_INTEGRATION_PROPOSAL.md) and re-run:

    python3 -m pip install --user python-docx
    python3 docs/build_recotap_docx.py

Output: docs/Recotap_Integration_Proposal.docx
"""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

OUT = Path(__file__).resolve().parent / "Recotap_Integration_Proposal.docx"

NAVY = RGBColor(0x1F, 0x2A, 0x37)
BLUE = RGBColor(0x17, 0x50, 0x89)
GREY = RGBColor(0x47, 0x55, 0x69)
CODE_FILL = "F4F6F8"
HEAD_FILL = "EAF0F7"


def _shade(element, fill):
    """Apply a solid background fill to a cell/paragraph XML element."""
    pr = element.get_or_add_tcPr() if element.tag.endswith("}tc") else element.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    pr.append(shd)


def add_heading(doc, text, level):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = NAVY if level == 1 else BLUE
    run.font.size = Pt(16 if level == 1 else 12.5)
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(4)


def add_para(doc, text, color=None, italic=False, size=10.5):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(6)
    return p


def add_bullets(doc, items, numbered=False):
    style = "List Number" if numbered else "List Bullet"
    for item in items:
        p = doc.add_paragraph(style=style)
        run = p.add_run(item)
        run.font.size = Pt(10.5)
        p.paragraph_format.space_after = Pt(2)


def add_code(doc, code):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.cell(0, 0)
    _shade(cell._tc, CODE_FILL)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    for i, line in enumerate(code.strip("\n").split("\n")):
        if i:
            p.add_run().add_break()
        r = p.add_run(line)
        r.font.name = "Consolas"
        r.font.size = Pt(8.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_table(doc, headers, rows):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    for i, h in enumerate(headers):
        _shade(hdr[i]._tc, HEAD_FILL)
        run = hdr[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
    for row in rows:
        cells = tbl.add_row().cells
        for i, val in enumerate(row):
            run = cells[i].paragraphs[0].add_run(val)
            run.font.size = Pt(9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def build():
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10.5)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = title.add_run("Beacon ↔ Recotap — Integration Proposal")
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = NAVY
    meta = doc.add_paragraph()
    mr = meta.add_run("From: Beacon (GTM / CRM platform)    •    For: Recotap")
    mr.font.size = Pt(10)
    mr.font.color.rgb = GREY
    add_para(doc, "Purpose: confirm scope and get your sign-off for a bi-directional API "
                  "integration between Beacon CRM and Recotap.", color=GREY, italic=True)

    add_heading(doc, "1. Summary", 1)
    add_para(doc, "Beacon is a GTM/CRM platform. We would like to connect to Recotap so that:")
    add_bullets(doc, [
        "Beacon → Recotap: we send our accounts, deals, and call/email activities so "
        "Recotap has live CRM context.",
        "Recotap → Beacon: we read back each account’s score, buying-journey stage, "
        "and intent sub-scores to surface inside Beacon.",
    ])
    add_para(doc, "Recotap remains the source of truth for intent/scoring; Beacon remains the "
                  "source of truth for CRM data. The integration is additive and read-only on "
                  "your scores — we never modify Recotap’s scoring logic or data.")

    add_heading(doc, "2. Data flow", 1)
    add_table(doc, ["Direction", "Data", "Source of truth"], [
        ["Beacon → Recotap", "Accounts (name + domain), Deals, Call/Email activities", "Beacon"],
        ["Recotap → Beacon", "Account score, journey stage, intent sub-scores", "Recotap"],
    ])
    add_para(doc, "Account match key: domain (lowercased, www-stripped). We also send our own "
                  "externalId for reference.")

    add_heading(doc, "3. API usage", 1)
    add_table(doc, ["Item", "Detail"], [
        ["Auth", "X-Api-Key header (per-workspace key you issue us)"],
        ["Transport", "HTTPS, JSON REST"],
        ["Push endpoints", "POST /accounts, /accounts/map-external-ids, /deals, /sales-activities, /deal-stages"],
        ["Pull endpoint", "GET /accounts (delta + pagination)"],
        ["Pagination", "keyset: cursor/limit → nextCursor, hasNextPage, syncTimestamp"],
        ["Delta sync", "lastSync=<ISO-8601> returns modified-since records only"],
        ["Batch sizes", "≤ 100 accounts · ≤ 100 deals ≤ 50 activities per request"],
        ["Webhooks", "Not required — we poll on a schedule"],
    ])

    add_heading(doc, "4. Example calls", 1)
    add_para(doc, "Illustrative payloads (exact field names per your API spec — happy to align). "
                  "Authentication header on every request:")
    add_code(doc, """X-Api-Key: <workspace_api_key>
Content-Type: application/json""")

    add_heading(doc, "4.1  Push accounts", 2)
    add_code(doc, """POST /api/v1/accounts
{
  "accounts": [
    {
      "externalId": "bcn_8f2a91",
      "name": "Haily HR",
      "domain": "hailyhr.com",
      "customFields": {
        "icp_score": 82,
        "account_thesis": "Mid-market HR SaaS, expanding PS team",
        "beacon_angle": "Implementation automation"
      }
    }
  ]
}""")

    add_heading(doc, "4.2  Push a deal (associated to an account by domain)", 2)
    add_code(doc, """POST /api/v1/deals
{
  "deals": [
    {
      "externalDealId": "bcn_deal_4471",
      "amount": 48000,
      "stageLabel": "Discovery",
      "stageId": "stg_discovery",
      "closedDate": "2026-08-15",
      "ownerEmail": "rep@beacon.li",
      "associatedAccounts": [{ "domain": "hailyhr.com" }]
    }
  ]
}""")

    add_heading(doc, "4.3  Push a call/email activity", 2)
    add_code(doc, """POST /api/v1/sales-activities
{
  "activities": [
    {
      "accountDomain": "hailyhr.com",
      "activityType": "call",
      "occurredAt": "2026-06-01T14:32:00Z",
      "durationMinutes": 12,
      "ownerEmail": "rep@beacon.li"
    }
  ]
}""")

    add_heading(doc, "4.4  Pull scores (delta + pagination)", 2)
    add_code(doc, """GET /api/v1/accounts?lastSync=2026-06-01T00:00:00Z&limit=100&cursor=<nextCursor>

200 OK
{
  "results": [
    {
      "domain": "hailyhr.com",
      "externalId": "bcn_8f2a91",
      "accountScore": 76,
      "journeyStage": "Consideration",
      "intentScores": { "linkedin": 71, "g2": 64, "bombora": 80, "website": 58 }
    }
  ],
  "nextCursor": "eyJpZCI6IDEwMjN9",
  "hasNextPage": true,
  "syncTimestamp": "2026-06-02T09:40:00Z"
}""")

    add_heading(doc, "5. Sync cadence & expected load", 1)
    add_para(doc, "Poll-based, delta-only:")
    add_bullets(doc, [
        "Push — about every 5 minutes; only records changed since last sync, batched within your limits.",
        "Pull — about every 15 minutes; GET /accounts?lastSync=…, paginated.",
    ])
    add_para(doc, "Steady-state volume is low (only changed records per cycle). The one-time "
                  "initial backfill is chunked within the batch limits above. We will honor any "
                  "rate limits you specify.")

    add_heading(doc, "6. What we send / read", 1)
    add_bullets(doc, [
        "Accounts: company name, domain, and a small set of CRM fields (e.g. ICP score, notes) as customFields.",
        "Deals: amount, stage, close date, owner email, associated account domain.",
        "Activities: calls and emails only — type, timestamp, duration. No meeting transcripts "
        "or notes unless you support and want them.",
        "Read (per account): score, journey stage, intent sub-scores — displayed in Beacon, strictly read-only.",
    ])

    add_heading(doc, "7. Security & data handling", 1)
    add_bullets(doc, [
        "API key stored encrypted per workspace; never exposed in any client/browser.",
        "All traffic over HTTPS.",
        "We send only B2B account and activity metadata — no consumer PII.",
        "Integration is feature-flagged and off by default, enabled only after both sides validate in a test workspace.",
    ])

    add_heading(doc, "8. What we need from Recotap", 1)
    add_bullets(doc, [
        "A sandbox / test API key + workspace to validate against.",
        "Confirm the HTTPS base URL and current endpoint paths.",
        "Rate limits and any concurrency constraints.",
        "Confirm supported activity types (we plan to send call + email).",
        "How to register deal stages — we will push our stage taxonomy via POST /deal-stages and reuse the IDs you return.",
        "Typical latency from an account push to its score being populated.",
        "Any pricing / plan implications for API access.",
    ], numbered=True)

    add_heading(doc, "9. Rollout", 1)
    add_bullets(doc, [
        "Validate in a test workspace with the sandbox key.",
        "One-time backfill: accounts → dedup map (map-external-ids) → deals → call/email activities → first full pull.",
        "Enable incremental delta sync (5-min push / 15-min pull) behind the flag.",
    ], numbered=True)

    add_para(doc, "Questions or changes to scope are welcome — we can adjust endpoints, "
                  "cadence, or field mapping to fit your platform’s constraints.",
             color=GREY, italic=True)

    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    build()
